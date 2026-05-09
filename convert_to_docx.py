"""
NVIDIA NPN 파트너 관리 기획서 v2 MD → DOCX 변환 스크립트
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MD_FILE = "/Users/sangokh/Desktop/기획서/NVIDIA_NPN_파트너_관리_기획서_v2.md"
OUT_FILE = "/Users/sangokh/Desktop/기획서/NVIDIA_NPN_파트너_관리_기획서_v2.docx"

# ── 색상 ──────────────────────────────────────────────────────────────────────
COLOR_NVIDIA_GREEN  = RGBColor(0x76, 0xB9, 0x00)   # #76B900
COLOR_DARK_GRAY     = RGBColor(0x33, 0x33, 0x33)
COLOR_LIGHT_GRAY_BG = "F2F2F2"
COLOR_HEADER_BG     = "1A1A2E"   # 짙은 네이비
COLOR_SUBHEADER_BG  = "D5E8D4"   # 연초록
COLOR_WHITE         = "FFFFFF"

FONT_DEFAULT = "맑은 고딕"


# ── 헬퍼 ──────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, sides=("top","bottom","left","right"), size=4, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_para(cell, text, bold=False, font_size=9, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    run  = para.add_run(text)
    run.font.name      = FONT_DEFAULT
    run.font.size      = Pt(font_size)
    run.font.bold      = bold
    if color:
        run.font.color.rgb = color
    return para


def add_styled_table(doc, headers, rows, col_widths_cm,
                     header_bg=COLOR_HEADER_BG, header_fg=COLOR_WHITE):
    """헤더 행 + 데이터 행이 있는 스타일 테이블 추가"""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 컬럼 너비
    for i, row_cells in enumerate(table.rows):
        for j, cell in enumerate(row_cells.cells):
            cell.width = Cm(col_widths_cm[j % len(col_widths_cm)])

    # 헤더 행
    hdr_cells = table.rows[0].cells
    for j, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, color="999999")
        cell_para(cell, hdr, bold=True, font_size=9,
                  color=RGBColor(0xFF,0xFF,0xFF) if header_bg == COLOR_HEADER_BG
                        else COLOR_DARK_GRAY)

    # 데이터 행
    for i, row_data in enumerate(rows):
        data_cells = table.rows[i + 1].cells
        bg = "FAFAFA" if i % 2 == 0 else COLOR_WHITE
        for j, (cell, text) in enumerate(zip(data_cells, row_data)):
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="DDDDDD")
            is_bold = (j == 0 and i == len(rows) - 1 and text.startswith("합계"))
            cell_para(cell, str(text), bold=is_bold, font_size=9)

    doc.add_paragraph()
    return table


def add_heading(doc, text, level):
    """헤딩 추가 (레벨 1~4)"""
    para = doc.add_heading(text, level=level)
    run  = para.runs[0] if para.runs else para.add_run(text)
    run.font.name = FONT_DEFAULT
    if level == 1:
        run.font.size  = Pt(18)
        run.font.color.rgb = RGBColor(0x76, 0xB9, 0x00)
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    elif level == 3:
        run.font.size  = Pt(12)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    else:
        run.font.size  = Pt(11)
        run.font.color.rgb = COLOR_DARK_GRAY
    para.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    para.paragraph_format.space_after  = Pt(4)
    return para


def add_body(doc, text, bullet=False, numbered=False, indent=0):
    """본문 단락 추가"""
    style = "List Bullet" if bullet else ("List Number" if numbered else "Normal")
    try:
        para = doc.add_paragraph(style=style)
    except Exception:
        para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    if indent:
        para.paragraph_format.left_indent = Cm(indent * 0.5)

    # 인라인 볼드 처리 (**text**)
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for k, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.name = FONT_DEFAULT
        run.font.size = Pt(10)
        run.font.bold = (k % 2 == 1)
    return para


def add_code_block(doc, text):
    """코드 블록 (회색 배경, 고정폭)"""
    for line in text.split("\n"):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent  = Cm(0.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run = para.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        # 배경색
        pPr = para._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "EBEBEB")
        pPr.append(shd)
    doc.add_paragraph()


# ── 마크다운 테이블 파서 ──────────────────────────────────────────────────────
def parse_md_table(lines):
    """마크다운 테이블 라인 목록 → (headers, rows)"""
    rows = []
    for line in lines:
        if re.match(r"^\s*\|[-: |]+\|\s*$", line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return [], []
    return rows[0], rows[1:]


# ── 표지 추가 ─────────────────────────────────────────────────────────────────
def add_cover(doc):
    # 녹색 배경 타이틀 박스 흉내 (border bottom)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(80)
    title_para.paragraph_format.space_after  = Pt(6)
    run = title_para.add_run("NVIDIA NPN")
    run.font.name  = FONT_DEFAULT
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = COLOR_NVIDIA_GREEN

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(4)
    run2 = sub_para.add_run("파트너 관리 기획서")
    run2.font.name  = FONT_DEFAULT
    run2.font.size  = Pt(22)
    run2.font.bold  = True
    run2.font.color.rgb = COLOR_DARK_GRAY

    # 구분선
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(4)
    sep.paragraph_format.space_after  = Pt(4)
    pPr = sep._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "76B900")
    pb.append(bottom)
    pPr.append(pb)

    meta_items = [
        ("작성일",   "2026-05-09"),
        ("버전",     "v2.0"),
        ("적용 범위", "한국 NPN 전체 — SP / SA / Distributor"),
        ("활용 목적", "내부 전략·운영 / 파트너 QBR 공유"),
    ]
    doc.add_paragraph()
    for label, val in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(f"{label}: ")
        r1.font.name  = FONT_DEFAULT
        r1.font.size  = Pt(10)
        r1.font.bold  = True
        r1.font.color.rgb = COLOR_DARK_GRAY
        r2 = p.add_run(val)
        r2.font.name  = FONT_DEFAULT
        r2.font.size  = Pt(10)
        r2.font.color.rgb = COLOR_DARK_GRAY

    # 페이지 나누기
    doc.add_page_break()


# ── 메인 변환 로직 ────────────────────────────────────────────────────────────
def convert(md_path, out_path):
    doc = Document()

    # 기본 스타일
    style = doc.styles["Normal"]
    style.font.name = FONT_DEFAULT
    style.font.size = Pt(10)

    # 여백
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    add_cover(doc)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # ── 헤딩 ──────────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.+)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            add_heading(doc, text, level)
            i += 1
            continue

        # ── 코드 블록 ─────────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # 닫는 ```
            add_code_block(doc, "\n".join(code_lines))
            continue

        # ── 마크다운 테이블 ───────────────────────────────────────────────────
        if line.strip().startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].rstrip("\n"))
                i += 1
            headers, rows = parse_md_table(tbl_lines)
            if headers:
                n = len(headers)
                base_w  = 15.5 / n
                col_w   = [base_w] * n
                # 첫 열 약간 넓게
                if n >= 3:
                    col_w[0] = base_w * 1.4
                    rest = (15.5 - col_w[0]) / (n - 1)
                    col_w[1:] = [rest] * (n - 1)
                add_styled_table(doc, headers, rows, col_w)
            continue

        # ── 수평선 ────────────────────────────────────────────────────────────
        if re.match(r"^---+$", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pb  = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "CCCCCC")
            pb.append(bot)
            pPr.append(pb)
            i += 1
            continue

        # ── 빈 줄 ─────────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── 인용 (> ...) ──────────────────────────────────────────────────────
        if line.strip().startswith("> "):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.8)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(text)
            run.font.name      = FONT_DEFAULT
            run.font.size      = Pt(9)
            run.font.italic    = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # ── 불릿 목록 ─────────────────────────────────────────────────────────
        if re.match(r"^\s*[-*]\s+", line):
            text   = re.sub(r"^\s*[-*]\s+", "", line)
            indent = (len(line) - len(line.lstrip())) // 2
            add_body(doc, text, bullet=True, indent=indent)
            i += 1
            continue

        # ── 번호 목록 ─────────────────────────────────────────────────────────
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            add_body(doc, text, numbered=True)
            i += 1
            continue

        # ── 일반 본문 ─────────────────────────────────────────────────────────
        add_body(doc, line.strip())
        i += 1

    doc.save(out_path)
    print(f"✅ 저장 완료: {out_path}")


if __name__ == "__main__":
    convert(MD_FILE, OUT_FILE)
